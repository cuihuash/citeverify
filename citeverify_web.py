"""Portable local HTML interface for the citation verifier.

Run this file locally, then open http://127.0.0.1:8765 in a browser. API keys
are read from command-line-specified local files and never placed in reports.
"""

from __future__ import annotations

import argparse
import html
import json
import re
import tempfile
import threading
import webbrowser
from dataclasses import asdict
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, urlparse

try:
    from email.parser import BytesParser
    from email.policy import default as email_policy
except ImportError:  # pragma: no cover
    BytesParser = None
    email_policy = None

from docx import Document

import citeverify_lookup as lookup
import citeverify_parser as citation_parser


MAX_UPLOAD_BYTES = 40 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc"}


def page_shell(title: str, body: str) -> str:
    return f"""<!doctype html>
<html lang="en"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{html.escape(title)}</title>
<style>
:root {{ color-scheme: light; font-family: Arial, sans-serif; color: #172033; background: #f4f6f8; }}
body {{ margin: 0; }}
main {{ max-width: 1080px; margin: 0 auto; padding: 36px 22px 60px; }}
.hero {{ background: #17324d; color: white; border-radius: 18px; padding: 30px; margin-bottom: 22px; }}
.hero h1 {{ margin: 0 0 10px; font-size: clamp(1.8rem, 4vw, 2.8rem); }}
.hero p {{ margin: 0; color: #dbeafe; max-width: 760px; }}
.panel {{ background: white; border: 1px solid #d8dee8; border-radius: 14px; padding: 24px; margin: 18px 0; box-shadow: 0 3px 12px #1720330d; }}
.dropzone {{ border: 2px dashed #8aa5c1; border-radius: 12px; padding: 30px; text-align: center; background: #f8fbff; }}
input[type=file] {{ display: block; margin: 16px auto; max-width: 100%; }}
button, .button {{ border: 0; border-radius: 8px; background: #0f766e; color: white; padding: 12px 18px; font-size: 1rem; cursor: pointer; text-decoration: none; display: inline-block; }}
button:hover, .button:hover {{ background: #115e59; }}
.muted {{ color: #5b6473; }}
.notice {{ border-left: 4px solid #d97706; background: #fffbeb; padding: 12px 15px; color: #78350f; }}
.error {{ border-left: 4px solid #dc2626; background: #fef2f2; padding: 12px 15px; color: #7f1d1d; }}
.file-list {{ text-align: left; max-width: 700px; margin: 12px auto; }}
.file-list li {{ margin: 5px 0; }}
.doc-section {{ margin-top: 36px; }}
.doc-section h2 {{ border-bottom: 2px solid #dbe3ed; padding-bottom: 10px; }}
.summary {{ display: flex; flex-wrap: wrap; gap: 8px; margin: 14px 0; }}
.summary-item, .status {{ display: inline-block; border-radius: 999px; padding: 5px 10px; font-weight: 700; font-size: .9rem; }}
.verified {{ background: #dcfce7; color: #166534; }}
.unverified {{ background: #fef3c7; color: #92400e; }}
.conflict, .unable-to-check {{ background: #fee2e2; color: #991b1b; }}
.verified-title-variant {{ background: #dbeafe; color: #1e40af; }}
.reference-card {{ background: white; border: 1px solid #d1d5db; border-radius: 10px; padding: 18px; margin: 14px 0; }}
.reference-heading {{ display: flex; align-items: center; justify-content: space-between; gap: 12px; }}
.reference-heading h3 {{ font-size: 1.05rem; margin: 0; }}
.citation {{ white-space: pre-wrap; overflow-wrap: anywhere; background: #f9fafb; border-left: 3px solid #9ca3af; padding: 10px 12px; }}
.processing {{ display: none; margin-top: 16px; border-left: 4px solid #2563eb; background: #eff6ff; padding: 12px 15px; color: #1e3a8a; }}
.processing.visible {{ display: block; }}
.selected-files {{ display: none; text-align: left; max-width: 700px; margin: 18px auto 0; border-top: 1px solid #d8dee8; padding-top: 14px; }}
.selected-files.visible {{ display: block; }}
.selected-files ul {{ margin: 8px 0 0; padding-left: 22px; }}
details {{ margin: 12px 0; }}
a {{ color: #075985; }}
</style></head><body><main>{body}</main></body></html>"""


def upload_page(error: str | None = None) -> str:
    error_html = f'<div class="error">{html.escape(error)}</div>' if error else ""
    body = f"""<section class="hero">
<h1>CiteVerify</h1>
<p>Upload one or more research papers. The app extracts the reference lists,
checks DOI and catalog information, and creates a readable HTML report.</p>
</section>
<section class="panel">
{error_html}
<form id="upload-form" method="post" action="/process" enctype="multipart/form-data">
<div class="dropzone">
<h2>Select your files</h2>
<p class="muted">Select multiple files at once, or choose files repeatedly to add them to the list.</p>
<input type="file" name="files" accept=".pdf,.docx,.doc" multiple required>
<div id="selected-files" class="selected-files" aria-live="polite">
  <strong id="selected-count">0 files selected</strong>
  <ul id="selected-file-list"></ul>
</div>
<button id="process-button" type="submit">Process references</button>
<div id="processing-message" class="processing" role="status" aria-live="polite">
  Processing has started. The app is extracting references and checking them against the lookup services. A paper with many references may take several minutes.
</div>
</div>
</form>
</section>
<section class="panel">
<h2>What the results mean</h2>
<p><strong>Verified</strong> means a source returned a strong match.</p>
<p><strong>Unverified</strong> means the app could not find enough evidence. It does not mean the citation is fabricated.</p>
<p class="notice">API keys stay on this computer. They are never included in the report.</p>
</section>"""
    body += """
<script>
const uploadForm = document.getElementById('upload-form');
const fileInput = document.querySelector('input[type=file]');
const selectedFiles = document.getElementById('selected-files');
const selectedCount = document.getElementById('selected-count');
const selectedFileList = document.getElementById('selected-file-list');
let chosenFiles = [];
if (fileInput) {
  fileInput.addEventListener('change', () => {
    const newlyChosenFiles = Array.from(fileInput.files || []);
    for (const file of newlyChosenFiles) {
      const alreadyChosen = chosenFiles.some(existing =>
        existing.name === file.name &&
        existing.size === file.size &&
        existing.lastModified === file.lastModified
      );
      if (!alreadyChosen) chosenFiles.push(file);
    }
    if (typeof DataTransfer !== 'undefined') {
      const transfer = new DataTransfer();
      chosenFiles.forEach(file => transfer.items.add(file));
      fileInput.files = transfer.files;
    }
    const files = chosenFiles;
    if (selectedCount) {
      selectedCount.textContent = files.length + (files.length === 1 ? ' file selected' : ' files selected');
    }
    if (selectedFileList) {
      selectedFileList.innerHTML = files.map(file => '<li>' + file.name.replace(/[&<>"']/g, character => ({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[character])) + '</li>').join('');
    }
    if (selectedFiles) selectedFiles.classList.toggle('visible', files.length > 0);
  });
}
if (uploadForm) {
  uploadForm.addEventListener('submit', () => {
    const button = document.getElementById('process-button');
    const message = document.getElementById('processing-message');
    if (button) {
      button.disabled = true;
      button.textContent = 'Processing…';
      button.style.opacity = '0.7';
      button.style.cursor = 'wait';
    }
    if (message) message.classList.add('visible');
  });
}
</script>"""
    return page_shell("CiteVerify", body)


def extract_docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_uploaded_file(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.casefold()
    if suffix not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {filename}")
    if suffix == ".doc":
        raise ValueError(f"{filename}: old .doc files are not supported; save it as .docx first.")

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as handle:
        handle.write(data)
        temporary_path = Path(handle.name)
    try:
        if suffix == ".pdf":
            return citation_parser.extract_pdf_text(temporary_path)
        return extract_docx_text(temporary_path)
    finally:
        temporary_path.unlink(missing_ok=True)


def parse_upload(body: bytes, content_type: str) -> list[tuple[str, bytes]]:
    if BytesParser is None or email_policy is None:
        raise RuntimeError("The Python email parser is unavailable.")
    header = f"Content-Type: {content_type}\r\nMIME-Version: 1.0\r\n\r\n".encode("utf-8")
    message = BytesParser(policy=email_policy).parsebytes(header + body)
    files: list[tuple[str, bytes]] = []
    for part in message.iter_attachments():
        filename = part.get_filename()
        if filename:
            files.append((Path(filename).name, part.get_payload(decode=True) or b""))
    return files


def process_files(
    files: list[tuple[str, bytes]],
    openalex_key: str | None,
    semantic_scholar_key: str | None,
) -> list[dict[str, Any]]:
    documents: list[dict[str, Any]] = []
    for filename, data in files:
        if len(data) > MAX_UPLOAD_BYTES:
            raise ValueError(f"{filename} is larger than the 40 MB upload limit.")
        text = extract_uploaded_file(filename, data)
        diagnostics = citation_parser.diagnose_text(text, filename)
        results = [
            lookup.verify_item(asdict(item), 0.2, openalex_key, semantic_scholar_key, index)
            for index, item in enumerate(diagnostics.diagnostics, start=1)
        ]
        documents.append({
            "filename": filename,
            "diagnostics": diagnostics,
            "results": results,
        })
    return documents


def report_body(documents: list[dict[str, Any]]) -> str:
    document_links: list[str] = []
    for index, document in enumerate(documents, start=1):
        anchor = f"document-{index}"
        filename = html.escape(str(document["filename"]))
        document_links.append(f'<li><a href="#{anchor}">{filename}</a></li>')

    navigation = ""
    if len(documents) > 1:
        navigation = (
            '<section class="panel" id="document-list">'
            '<h2>Jump to a document</h2>'
            '<p class="muted">Select a file to jump directly to its results.</p>'
            f'<ol class="file-list">{"".join(document_links)}</ol>'
            '</section>'
        )

    sections: list[str] = [
        '<section class="hero" id="results-top"><h1>CiteVerify results</h1>'
        f'<p>{len(documents)} document{"s" if len(documents) != 1 else ""} processed. '
        'Conflicts appear first, followed by unverified items and then verified items. '
        'Unverified does not mean fabricated.</p></section>',
        navigation,
    ]
    for index, document in enumerate(documents, start=1):
        report = lookup.render_html_report(document["filename"], document["results"])
        body_match = re.search(r"<body>(.*)</body>", report, flags=re.S)
        inner = body_match.group(1) if body_match else report
        inner = inner.replace('<h1>CiteVerify report</h1>', '')
        filename = html.escape(str(document["filename"]))
        diagnostics = document.get("diagnostics")
        if not document["results"]:
            references_header = html.escape(
                getattr(diagnostics, "references_header", None) or "not detected"
            )
            warnings = getattr(diagnostics, "warnings", []) or []
            warning_html = "".join(
                f"<li>{html.escape(str(warning))}</li>" for warning in warnings
            )
            if not warning_html:
                warning_html = "<li>No parser warning was recorded.</li>"
            inner += (
                '<div class="notice">'
                '<strong>No references were extracted from this document.</strong>'
                '<p>This usually means the PDF reference heading was not recognized, '
                'the reference list is an image scan, or the PDF layout needs a parser adjustment.</p>'
                f"<p><strong>References heading detected:</strong> {references_header}</p>"
                f"<p><strong>Parser notes:</strong></p><ul>{warning_html}</ul>"
                '</div>'
            )
        back_target = "#document-list" if len(documents) > 1 else "#results-top"
        sections.append(
            f'<section class="doc-section" id="document-{index}">'
            f'<h2>{filename}</h2>{inner}'
            f'<p><a href="{back_target}">Back to document list</a></p>'
            '</section>'
        )
    sections.append('<p><a class="button" href="/">Process another set of files</a></p>')
    return "".join(sections)


class CitationHandler(BaseHTTPRequestHandler):
    server_version = "CiteVerify/0.1.0"

    def do_GET(self) -> None:  # noqa: N802
        path = urlparse(self.path).path
        if path == "/favicon.ico":
            self.send_response(HTTPStatus.NO_CONTENT)
            self.end_headers()
            return
        if path != "/":
            self.send_error(HTTPStatus.NOT_FOUND)
            return
        self.send_html(upload_page())

    def do_POST(self) -> None:  # noqa: N802
        if urlparse(self.path).path != "/process":
            self.send_error(HTTPStatus.NOT_FOUND)
            return
        try:
            length = int(self.headers.get("Content-Length", "0"))
            if length <= 0 or length > MAX_UPLOAD_BYTES * 10:
                raise ValueError("The upload is empty or exceeds the total upload limit.")
            content_type = self.headers.get("Content-Type", "")
            files = parse_upload(self.rfile.read(length), content_type)
            if not files:
                raise ValueError("No files were selected.")
            documents = process_files(files, self.server.openalex_key, self.server.semantic_scholar_key)
            self.send_html(page_shell("CiteVerify results", report_body(documents)))
        except Exception as exc:
            self.send_html(upload_page(str(exc)), status=HTTPStatus.BAD_REQUEST)

    def send_html(self, content: str, status: HTTPStatus = HTTPStatus.OK) -> None:
        encoded = content.encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.send_header("Content-Length", str(len(encoded)))
        self.end_headers()
        self.wfile.write(encoded)

    def log_message(self, format: str, *args: Any) -> None:
        print(f"[{self.log_date_time_string()}] {format % args}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Run the local citation HTML interface.")
    parser.add_argument("--host", default="127.0.0.1")
    parser.add_argument("--port", type=int, default=8765)
    parser.add_argument("--openalex-key-file", type=Path)
    parser.add_argument("--s2-api-key-file", type=Path)
    parser.add_argument("--no-browser", action="store_true", help="Do not open the browser automatically")
    args = parser.parse_args()

    openalex_key = lookup.read_api_key_file(args.openalex_key_file) if args.openalex_key_file else None
    semantic_scholar_key = lookup.read_api_key_file(args.s2_api_key_file) if args.s2_api_key_file else None
    server = ThreadingHTTPServer((args.host, args.port), CitationHandler)
    server.openalex_key = openalex_key
    server.semantic_scholar_key = semantic_scholar_key
    url = f"http://{args.host}:{args.port}/"
    print(f"CiteVerify is running at {url}")
    if not args.no_browser:
        threading.Timer(0.5, lambda: webbrowser.open(url)).start()
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nStopping CiteVerify.")
    finally:
        server.server_close()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
